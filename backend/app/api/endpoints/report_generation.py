"""
Report Generation Endpoint - 测试报告数据整理
"""
from fastapi import APIRouter, UploadFile, File, HTTPException
from typing import List
import uuid

from app.models.schemas import (
    ReportGenerationRequest,
    ReportGenerationResult,
    UploadResponse
)
from app.services.file_service import FileService
from app.services.job_service import JobService

router = APIRouter()
file_service = FileService()
job_service = JobService()


@router.post("/upload-template", response_model=UploadResponse)
async def upload_report_template(file: UploadFile = File(...)):
    """
    上传Word报告模板

    文件要求:
    - 格式: .docx
    - 包含需要填充的表格
    """
    if not file.filename.endswith('.docx'):
        raise HTTPException(status_code=400, detail="仅支持Word文档 (.docx)")

    file_id = str(uuid.uuid4())
    saved_path = await file_service.save_upload(file, file_id)

    return UploadResponse(
        file_id=file_id,
        filename=file.filename,
        size=file.size or 0,
        message="报告模板上传成功"
    )


@router.post("/upload-data", response_model=UploadResponse)
async def upload_report_data(file: UploadFile = File(...)):
    """
    上传报告数据Excel文件

    文件要求:
    - 格式: .xlsx
    - 包含测试数据 (HTRB/H3TRB/HTGB/TC/IOL等)
    """
    if not file.filename.endswith(('.xlsx', '.xls')):
        raise HTTPException(status_code=400, detail="仅支持Excel文件 (.xlsx, .xls)")

    file_id = str(uuid.uuid4())
    saved_path = await file_service.save_upload(file, file_id)

    return UploadResponse(
        file_id=file_id,
        filename=file.filename,
        size=file.size or 0,
        message="数据文件上传成功"
    )


@router.post("/generate", response_model=dict)
async def generate_report(request: ReportGenerationRequest):
    """
    生成测试报告

    处理流程:
    1. 加载Word模板
    2. 读取Excel测试数据
    3. 匹配表格与数据
    4. 填充VF/IR/BV统计数据
    5. 更新变化率信息
    6. 保存新报告

    支持的报告类型:
    - HTRB: 高温反偏
    - H3TRB: 高温高湿
    - HTGB+: 高温正向栅偏
    - HTGB-: 高温负向栅偏
    - TC: 温度循环
    - IOL: 闭锁电流
    - AC: 交流
    """
    job_id = await job_service.create_job(
        job_type="report_generation",
        params=request.model_dump()
    )

    from app.workers.tasks.processing_tasks import process_report_generation_task
    process_report_generation_task.delay(job_id, request.model_dump())

    return {
        "job_id": job_id,
        "status": "pending",
        "message": "报告生成任务已创建"
    }


@router.get("/report-types")
async def get_report_types():
    """
    获取支持的报告类型列表
    """
    return {
        "report_types": [
            {
                "code": "HTRB",
                "name": "高温反偏",
                "description": "High Temperature Reverse Bias",
                "duration": "1000h"
            },
            {
                "code": "H3TRB",
                "name": "高温高湿",
                "description": "High Humidity High Temperature Reverse Bias",
                "duration": "1000h"
            },
            {
                "code": "HTGB+",
                "name": "高温正向栅偏",
                "description": "High Temperature Gate Bias (Positive)",
                "duration": "1000h"
            },
            {
                "code": "HTGB-",
                "name": "高温负向栅偏",
                "description": "High Temperature Gate Bias (Negative)",
                "duration": "1000h"
            },
            {
                "code": "TC",
                "name": "温度循环",
                "description": "Temperature Cycling",
                "duration": "1000h"
            },
            {
                "code": "IOL",
                "name": "闭锁电流",
                "description": "Latch-up Current Test",
                "duration": "96h"
            },
            {
                "code": "AC",
                "name": "交流测试",
                "description": "AC Test",
                "duration": "N/A"
            }
        ]
    }


@router.get("/preview/{file_id}")
async def preview_report_template(file_id: str):
    """
    预览报告模板结构

    返回:
    - 表格列表
    - 需要填充的字段
    """
    file_info = await file_service.get_file_info(file_id)
    if not file_info:
        raise HTTPException(status_code=404, detail="文件不存在")

    from docx import Document
    doc = Document(file_info["path"])

    tables_info = []
    for i, table in enumerate(doc.tables):
        # 获取表格标题 (通常在表格前一段)
        rows_count = len(table.rows)
        cols_count = len(table.columns) if table.rows else 0

        tables_info.append({
            "table_index": i,
            "rows": rows_count,
            "columns": cols_count,
            "preview": [
                [cell.text[:50] for cell in row.cells]
                for row in table.rows[:3]  # 预览前3行
            ]
        })

    return {
        "file_id": file_id,
        "filename": file_info["filename"],
        "tables_count": len(doc.tables),
        "tables": tables_info
    }